from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION, WD_ORIENTATION

doc = Document()

# modify the section margins
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(0.5)

# add header and footer 
header = section.header
header.paragraphs[0].text = "Section 1 Header"
footer = section.footer
footer.paragraphs[0].text = "Section 1 Footer"  

# add a paragraph
para1 = doc.add_paragraph("This is the first section with modified margins, header, and footer.")
para1.runs[0].bold = True
para1.runs[0].font.size = Pt(12)

# add a new section with different orientation and margins 
section2 = doc.add_section(WD_SECTION.NEW_PAGE)

section2.orientation = WD_ORIENTATION.LANDSCAPE

section2.page_height, section2.page_width = section2.page_width, section2.page_height # mannually swap dimensions to effect oreintation change

# unlink headers and footers from previous section
section2.header.is_linked_to_previous = False
section2.footer.is_linked_to_previous = False

doc.add_paragraph("This is the second section in landscape orientation")

section2.header.paragraphs[0].text = "Section 2 Header"
section2.footer.paragraphs[0].text = "Section 2 Footer"

doc.save("sections.docx")