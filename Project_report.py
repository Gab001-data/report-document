import sys
import argparse
import logging
import re
import os
from datetime import datetime, date
from collections import defaultdict
import tempfile
from urllib.parse import urljoin
import json  # ensure you import this at top if not yet present!

import requests
import pandas as pd
from bs4 import BeautifulSoup

from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor, Pt, Inches, Mm


from pdf2docx import Converter

# ─── Logger Setup ───────────────────────────────────────────
def ensure_str(val):
    return ", ".join(val) if isinstance(val, (list, tuple)) else str(val or "N/A")

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# ─── API & Session ──────────────────────────────────────────
API_TOKEN = os.getenv("HIGHBOND_API_TOKEN", "e1fd78c5aa8f83cb1633699c0b05d7a881bc926210bf7bb2272b40f90167279b")
BASE_URL  = "https://apis-us.diligentoneplatform.com/v1/orgs/52734"
HEADERS   = {
    "Authorization": f"Bearer {API_TOKEN}",
    "Content-Type": "application/vnd.api+json"
}

session = requests.Session()
retries = Retry(total=5, backoff_factor=1, status_forcelist=[429,500,502,503,504], allowed_methods=["GET"])
adapter = HTTPAdapter(max_retries=retries)
session.mount("https://", adapter)
session.mount("http://", adapter)

def convert_pdf_to_docx(pdf_path: str, docx_path: str):
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        logger.info(f"✅ PDF converted to DOCX: {docx_path}")
    except Exception as e:
        logger.error(f"❌ Error in PDF→DOCX conversion: {e}")

# ─── Data Fetching ──────────────────────────────────────────
def get_all_projects():
    today_str = date.today().isoformat()
    url = f"{BASE_URL}/projects"
    params = {"filter[start_date][lte]": today_str, "page[size]": 100, "page[number]": 1}
    all_projects = []
    while url:
        try:
            resp = session.get(url, headers=HEADERS, params=params, timeout=10)
            resp.raise_for_status()
        except Exception as e:
            logger.error(f"Failed to fetch projects: {e}")
            break
        js = resp.json()
        all_projects.extend(js.get("data", []))
        next_link = js.get("links", {}).get("next")
        url = urljoin(BASE_URL, next_link) if next_link else None
        params = None
    # final date filter
    valid = []
    for p in all_projects:
        sd = p.get("attributes", {}).get("start_date", "")
        try:
            if datetime.strptime(sd, "%Y-%m-%d").date() <= date.today():
                valid.append(p)
        except ValueError:
            continue
    logger.info(f"📦 Retrieved {len(valid)} valid projects")
    return valid

def get_project_issues(pid):
    try:
        resp = session.get(f"{BASE_URL}/projects/{pid}/issues", headers=HEADERS, timeout=10)
        resp.raise_for_status()
    except Exception as e:
        logger.error(f"Failed to fetch issues for project {pid}: {e}")
        return []
    return resp.json().get("data", [])

def clean_html_and_extract_tables(value):
    if not isinstance(value, str):
        return ensure_str(value), []

    soup = BeautifulSoup(value, "html.parser")
    for tag in soup(["script", "style"]): tag.decompose()
    for br in soup.find_all("br"): br.replace_with("\n")
    for p in soup.find_all("p"): p.insert_after("\n")

    tables = []
    for tbl in soup.find_all("table"):
        headers = [th.get_text(strip=True) for th in tbl.find_all("th")]
        rows = [
            [td.get_text(strip=True) for td in tr.find_all("td")]
            for tr in tbl.find_all("tr")
            if any(td.get_text(strip=True) for td in tr.find_all("td"))
        ]
        tables.append((headers, rows))
        tbl.decompose()

    text = soup.get_text(separator="\n").strip()
    text = text.replace("$", "")
    #text = re.sub(r"[•–-]\s*", "", text)  # Remove bullets if needed
    text = re.sub(r"\n+", "\n", text)     # Normalize newlines

    return text or "N/A", tables



# Regex precompiled
pat1 = re.compile(r"Management\s*Comment\s*1[:\-–]?\s*(.*)",)
pat2 = re.compile(r"Management\s*Comment\s*2[:\-–]?\s*(.*)",)

def extract_from_text(text, which):
    """Regex + loose fallback"""
    if not isinstance(text, str): return ""
    text = text.strip()
    if which == 1:
        m = pat1.search(text)
        if m: return m.group(1).strip()
        if "Custom field 1" in text:
            idx = text.find("Management comment 1") + len("management comment 1")
            return text[idx:].strip()
    else:
        m = pat2.search(text)
        if m: return m.group(1).strip()
        if "Custom field 2" in text:
            idx = text.lower().find("Management comment 2") + len("management comment 2")
            return text[idx:].strip()
    return ""

def deep_find_comments(data):
    """
    Recursively search ALL nested dicts/lists for Management Comment 1/2.
    Return tuple (cm1, cm2).
    """
    pat1 = re.compile(r"Management\s*Comment\s*1[:\-–]?\s*(.+)",)
    pat2 = re.compile(r"Management\s*Comment\s*2[:\-–]?\s*(.+)",)

    cm1 = ""
    cm2 = ""

    def walk(obj):
        nonlocal cm1, cm2
        if isinstance(obj, dict):
            for k, v in obj.items():
                if isinstance(v, (dict, list)):
                    walk(v)
                elif isinstance(v, str):
                    text = BeautifulSoup(v, "html.parser").get_text(separator="\n").strip()
                    if not cm1:
                        m = pat1.search(text)
                        if m:
                            cm1 = m.group(1).strip()
                    if not cm2:
                        m = pat2.search(text)
                        if m:
                            cm2 = m.group(1).strip()
        elif isinstance(obj, list):
            for item in obj:
                walk(item)

    walk(data)
    return cm1, cm2

# ─── Width Calculator ────────────────────────────────────────
def _compute_column_widths(text_matrix, max_total_width_inches=10.0, min_width_inches=0.5):
    if not text_matrix:
        return []
    cols = list(zip(*text_matrix))
    scores, total = [], 0.0
    for col in cols:
        lengths = [len(str(c)) for c in col]
        mx = max(lengths)
        has_sent = any(len(str(c).split()) > 5 for c in col)
        is_num = all(str(c).replace(".", "", 1).isdigit() for c in col if c)
        weight = 1.5 if has_sent else 0.5 if is_num else 1.0
        scores.append(mx * weight)
        total += mx * weight
    total = total or 1.0
    widths = [(s / total) * max_total_width_inches for s in scores]
    widths = [max(min_width_inches, w) for w in widths]
    used = sum(widths)
    if used > max_total_width_inches:
        factor = max_total_width_inches / used
        widths = [w * factor for w in widths]
    return [Inches(w) for w in widths]


# ─── Mini-Table Insertion ────────────────────────────────────
def add_mini_table_to_cell(cell, headers, rows):
    # flatten into a matrix
    col_count = max([len(headers)] + [len(r) for r in rows] + [0])
    matrix = []
    if headers:
        matrix.append([headers[i] if i < len(headers) else "" for i in range(col_count)])
    for r in rows:
        matrix.append([r[i] if i < len(r) else "" for i in range(col_count)])

    # Compute column widths
    max_w = getattr(cell, "width", Inches(5)).inches
    col_w = _compute_column_widths(matrix, max_total_width_inches=max_w)

    # Create the mini table
    mini = cell.add_table(rows=1 if headers else 0, cols=col_count)
    mini.autofit = False
    mini.style = "Light Grid Accent 1"

    if headers:
        hdr_row = mini.rows[0].cells
        for idx, txt in enumerate(matrix[0]):
            hdr_row[idx].width = col_w[idx]
            p = hdr_row[idx].paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)

    for row_vals in matrix[1 if headers else 0:]:
        row_cells = mini.add_row().cells
        for idx, txt in enumerate(row_vals):
            row_cells[idx].width = col_w[idx]
            p = row_cells[idx].paragraphs[0]
            p.text = txt
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15

    return mini  # ✅ Always return the created mini-table

# ─── Build Final DOCX ───────────────────────────────────────
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DARK_BLUE = RGBColor(0, 51, 102)
ORANGE_HEX = "FFA500"
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Mm
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
from datetime import date

from collections import defaultdict
from datetime import date
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Mm
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Make sure you have these functions defined elsewhere:
# def clean_html_and_extract_tables(html): ...
# def add_mini_table_to_cell(cell, headers, rows): ...

from collections import defaultdict
from datetime import date
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Mm
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ➜ Make sure these exist:
# def clean_html_and_extract_tables(html): ...
# def add_mini_table_to_cell(cell, headers, rows): ...

def normalize_region(region):
    if isinstance(region, str):
        return region.strip().title()
    elif isinstance(region, list):
        return [r.strip().title() for r in region]
    else:
        return str(region)

def add_footer_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.add_paragraph()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run = p.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_word_report(table_data, region_filters, severity_list):
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)

    # ── Styles ──
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)

    # ── Normalize regions ──
    normalized_regions = normalize_region(region_filters) if region_filters else ["All Regions"]
    regions_display = ", ".join(normalized_regions) if isinstance(normalized_regions, list) else normalized_regions

    # ── Find earliest start date ──
    earliest_start = None
    for row in table_data:
        if len(row) > 4:
            try:
                dt = datetime.strptime(row[4], "%Y-%m-%d")
                if earliest_start is None or dt < earliest_start:
                    earliest_start = dt
            except:
                continue

    # fallback if not found
    if earliest_start is None:
        earliest_start = date.today()
    month_str = earliest_start.strftime("%Y-%m")
    start_date_str = earliest_start.strftime("%Y-%m-%d")

    # ── COVER PAGE ──
    cover_table = doc.add_table(rows=2, cols=3)
    cover_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_table.autofit = False
    cover_table.columns[0].width = Inches(1.5)
    cover_table.columns[1].width = Inches(8)
    cover_table.columns[2].width = Inches(1.5)

    try:
        para_l = cover_table.cell(0, 0).paragraphs[0]
        para_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_l.add_run().add_picture("images.png", width=Inches(1.4))
    except Exception as e:
        print(f"⚠️ Left logo: {e}")

    para_c = cover_table.cell(0, 1).paragraphs[0]
    para_c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        para_r = cover_table.cell(0, 2).paragraphs[0]
        para_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para_r.add_run().add_picture("minigroup_logo.png", width=Inches(1.4))
    except Exception as e:
        print(f"⚠️ Right logo: {e}")

    title_cell = cover_table.cell(1, 0).merge(cover_table.cell(1, 2))
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p.add_run(f"Regional Issues Report")
    title_run.font.size = Pt(32)
    title_run.font.name = "Calibri"
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph("Mini Group / Eleven Degrees Consulting").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    start_date_obj = datetime.strptime(start_date_str, "%Y-%m-%d")
    year_month_str = start_date_obj.strftime("%Y-%m")
    p_info.add_run(f"Region(s): {regions_display} | Month: {year_month_str}")

    doc.add_paragraph(f"Date: {date.today():%Y-%m-%d}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    doc.add_page_break()

    # ── Start new section for content ──
    doc.add_section(WD_SECTION.NEW_PAGE)
    content_section = doc.sections[-1]

    content_section.header.is_linked_to_previous = False
    content_section.footer.is_linked_to_previous = False

    # Restart page numbers
    sectPr = content_section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), "1")

    # Add running header with correct month + region only (no date)
    header_p = content_section.header.add_paragraph(f"{month_str} {regions_display} Region Issues Report")
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].italic = True

    add_footer_page_number(content_section)

    # ── Group by Project ID ──
    grouped = defaultdict(list)
    for row in table_data:
        grouped[row[0]].append(row)

    first = True
    for pid, rows in grouped.items():
        if not first:
            doc.add_section(WD_SECTION.NEW_PAGE)
            content_section = doc.sections[-1]
            content_section.header.is_linked_to_previous = False
            content_section.footer.is_linked_to_previous = False

            header_p = content_section.header.add_paragraph(f"{month_str} {regions_display} Region Issues Report")
            header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header_p.runs[0].italic = True

            add_footer_page_number(content_section)

        proj = rows[0]
        auditor = proj[17] if len(proj) > 17 else "N/A"
        name, branch, region, start, status = proj[1], proj[2], proj[3], proj[4], proj[5]
        bm, om, sup = proj[14], proj[15], proj[16]

        doc.add_heading(f"Project: {name}", level=1)
        for label, value in [
            ("Branch", branch),
            ("Branch Manager", bm),
            ("Auditor(s)", auditor)
        ]:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value or "N/A"))

        doc.add_paragraph()

        footer = content_section.footer
        for p in list(footer.paragraphs):
            footer._element.remove(p._element)

        footer_line = f"Branch: {branch} | Region: {region} | Start: {start} | BM: {bm} | OM: {om} | Sup: {sup} | Auditors: {auditor}"
        p = footer.add_paragraph(footer_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 51, 102)

        for issue in rows:
            issue_title = issue[6] or "Untitled Issue"
            doc.add_heading(f"Issue: {issue_title}", level=2)

            desc_text, desc_tables = clean_html_and_extract_tables(issue[8])
            impl_text, _ = clean_html_and_extract_tables(issue[9])
            rec_text, rec_tables = clean_html_and_extract_tables(issue[13])

            cost_impact = issue[10]
            if isinstance(cost_impact, str):
                cost_impact = cost_impact.replace("$", "").strip()
            try:
                cost_impact = float(cost_impact)
            except Exception:
                cost_impact = 0

            fields = [
                ("Severity", issue[7]),
                ("Description", desc_text),
                ("Implication", impl_text),
                ("Cost Impact", f"{cost_impact:,.2f}"),
                ("Management Comment 1", issue[11]),
                ("Management Comment 2", issue[12]),
                ("Recommendation", rec_text),
            ]

            tbl = doc.add_table(rows=len(fields), cols=2)
            tbl.style = "Table Grid"
            tbl.autofit = False
            tbl.allow_autofit = False

            tblPr = tbl._tbl.tblPr
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)

            tbl.columns[0].width = Inches(1.8)
            tbl.columns[1].width = Inches(8.2)

            set_table_border_color(tbl, 'FF9900')

            for i, (lbl, val) in enumerate(fields):
                c0, c1 = tbl.rows[i].cells
                c0.text = lbl
                c0.paragraphs[0].runs[0].bold = True

                p = c1.paragraphs[0]
                for run in p.runs:
                    p._p.remove(run._r)
                p.add_run(str(val or "N/A"))

                nested_tables = []
                if lbl == "Description":
                    nested_tables = desc_tables
                elif lbl == "Recommendation":
                    nested_tables = rec_tables

                for hdrs, rows_tbl in nested_tables:
                    mini_table = add_mini_table_to_cell(c1, hdrs, rows_tbl)
                    for hdr_cell in mini_table.rows[0].cells:
                        for para in hdr_cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0, 102, 204)

            doc.add_paragraph()

        first = False

    return doc









def update_footer(section, branch, region, start_date):
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing paragraphs
    for p in list(footer.paragraphs):
        footer._element.remove(p._element)

    # Custom line
    line = f"Branch: {branch}   |   Region: {region}   |   Start: {start_date}"
    p = footer.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 51, 102)


# ─── Deep JSON fallback extractor ─────────────────────────────
def extract_from_text(raw_json: str, idx: int) -> str:
    """
    Last-resort regex over the JSON dump to grab
    'Comment 1: text...' or 'Management Comment 2 – text...'
    """
    pat = re.compile(rf"(?:Management|Manager)?\s*Comment\s*{idx}\s*[:\-–]\s*(.+?)(?=[\"'\n\r]|\}})", re.IGNORECASE)
    m = pat.search(raw_json)
    return m.group(1).strip() if m else ""

# ─── Primary comment finder ────────────────────────────────────
def deep_find_comments(data):
    """
    Recursively walk any dict/list/str looking
    for your 1st/2nd Management Comments in TEXT.
    """
    # canonical lowercase keys to detect
    keys_1 = ["Custom field 1", "Management comment 1", "manager comment 1", "comment 1"]
    keys_2 = ["Custom field 2", "Management comment 2", "manager comment 2", "comment 2"]

    pat1 = re.compile(r"(?:Custom field 1)?\s*Comment\s*1\s*[:\-–]\s*(.+)",)
    pat2 = re.compile(r"(?:Management|Manager)?\s*Comment\s*2\s*[:\-–]\s*(.+)", re.IGNORECASE)

    cm1, cm2 = "", ""

    def walk(obj):
        nonlocal cm1, cm2
        if isinstance(obj, dict):
            for k, v in obj.items():
                k_l = str(k).strip().lower()

                # check attribute names (like custom_attributes entries)
                if not cm1 and any(k_l == key for key in keys_1) and isinstance(v, str):
                    cm1 = BeautifulSoup(v, "html.parser").get_text().strip()
                if not cm2 and any(k_l == key for key in keys_2) and isinstance(v, str):
                    cm2 = BeautifulSoup(v, "html.parser").get_text().strip()

                walk(v)

        elif isinstance(obj, list):
            for item in obj:
                walk(item)

        elif isinstance(obj, str):
            text = BeautifulSoup(obj, "html.parser").get_text().strip()

            if not cm1:
                m1 = pat1.search(text)
                if m1:
                    cm1 = m1.group(1).strip()

            if not cm2:
                m2 = pat2.search(text)
                if m2:
                    cm2 = m2.group(1).strip()

    walk(data)
    return cm1, cm2



# ─── Main & CLI ─────────────────────────────────────────────
# ─── Main ─────────────────────────────────────────────
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# ─── Deep JSON fallback extractor ─────────────────────────────
def extract_from_text(raw_json: str, idx: int) -> str:
    """
    Last-resort regex over the JSON dump to grab
    'Comment 1: text...' or 'Management Comment 2 – text...'
    """
    pat = re.compile(
        rf"(?:Management|Manager)?\s*Comment\s*{idx}\s*[:\-–]\s*(.+?)(?=[\"'\n\r]|\}})",
        re.IGNORECASE,
    )
    m = pat.search(raw_json)
    return m.group(1).strip() if m else ""

def deep_list_all_fields(obj, prefix=""):
    """
    Recursively list all field keys and values under a given dict/list.
    Useful for identifying all artificial/custom fields.
    """
    results = []

    if isinstance(obj, dict):
        for k, v in obj.items():
            full_key = f"{prefix}.{k}" if prefix else k
            if isinstance(v, (dict, list)):
                results.extend(deep_list_all_fields(v, prefix=full_key))
            else:
                results.append((full_key, v))
    elif isinstance(obj, list):
        for idx, item in enumerate(obj):
            full_key = f"{prefix}[{idx}]"
            results.extend(deep_list_all_fields(item, prefix=full_key))
    return results


pat1 = re.compile(r"(?:Management|Manager)?\s*Comment\s*1[:\-–]?\s*(.+)", re.IGNORECASE)
pat2 = re.compile(r"(?:Management|Manager)?\s*Comment\s*2[:\-–]?\s*(.+)", re.IGNORECASE)

# ─── UTILITIES ────────────────────────────────────────
def clean_html(txt):
    """Strip HTML tags, convert <br> to newline."""
    soup = BeautifulSoup(txt, "html.parser")
    for br in soup.find_all("br"): br.replace_with("\n")
    return soup.get_text().strip()

def convert_pdf_to_text(pdf_bytes: bytes) -> str:
    """Save PDF temp→DOCX→extract text via pdf2docx."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as fp_pdf:
        fp_pdf.write(pdf_bytes)
        fp_pdf.flush()
        docx_path = fp_pdf.name + ".docx"
        Converter(fp_pdf.name).convert(docx_path, start=0, end=None).close()
    # now read docx as plain text
    text = []
    from docx import Document
    for p in Document(docx_path).paragraphs:
        text.append(p.text)
    return "\n".join(text)

def deep_list_all_fields(obj, prefix=""):
    """
    Recursively list all field keys and scalar values.
    Returns list of (field_path, value).
    """
    results = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            path = f"{prefix}.{k}" if prefix else k
            if isinstance(v, (dict, list)):
                results += deep_list_all_fields(v, path)
            else:
                results.append((path, v))
    elif isinstance(obj, list):
        for idx, item in enumerate(obj):
            path = f"{prefix}[{idx}]"
            results += deep_list_all_fields(item, path)
    return results


def set_table_border_color(table, color='FF9900'):
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)

    tblPr.append(tblBorders)

# --------------------------
# Function: Set mini-table header color to blue
# --------------------------
def style_mini_table_header(mini_table):
    header_row = mini_table.rows[0]
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 102, 204)  # Blue

# --------------------------
# Function: Add footer with page number
# --------------------------
def add_footer_with_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False

    # Add custom footer text if you want
    p = footer.paragraphs[0]
    p.text = "Confidential Report"

    # Add page number field
    page_number_paragraph = footer.add_paragraph()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run = page_number_paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def fetch_issue_attachments(issue):
    """
    If an issue has attachments, download and return dict {filename: bytes}
    """
    atts = {}
    rel = issue.get("relationships", {}).get("attachments", {}).get("data", [])
    for meta in rel:
        att_id = meta.get("id")
        if not att_id: continue
        url = f"{BASE_URL}/attachments/{att_id}/content"
        logger.debug(f"Fetching attachment {att_id} → {url}")
        r = session.get(url, headers=HEADERS, timeout=15)
        if r.ok:
            disp = r.headers.get("Content-Disposition","")
            fn = re.search(r'filename="?(.*?)"?(;|$)', disp)
            name = fn.group(1) if fn else f"attachment_{att_id}"
            atts[name] = r.content
        else:
            logger.warning(f"Failed to download attachment {att_id}: {r.status_code}")
    return atts

def extract_comments_from_text(text):
    """Run both regexes on a block of text and return (cm1, cm2) or empty."""
    cm1 = cm2 = ""
    m1 = pat1.search(text)
    if m1: cm1 = m1.group(1).strip()
    m2 = pat2.search(text)
    if m2: cm2 = m2.group(1).strip()
    return cm1, cm2

def deep_find_comments(data):
    """
    Recursively search HTML/Text inside nested dicts/lists.
    Returns first found (cm1, cm2).
    """
    cm1 = cm2 = ""
    def walk(o):
        nonlocal cm1, cm2
        if isinstance(o, str):
            txt = clean_html(o)
            c1, c2 = extract_comments_from_text(txt)
            if c1 and not cm1: cm1 = c1
            if c2 and not cm2: cm2 = c2
        elif isinstance(o, dict):
            for v in o.values():
                walk(v)
        elif isinstance(o, list):
            for i in o:
                walk(i)
    walk(data)
    return cm1, cm2

# ─── CORE COMMENT EXTRACTION ────────────────────────────
def find_management_comments(issue, project_custom_attrs):
    """
    1) Gather ALL candidate texts from:
       • issue["custom_attributes"]
       • issue top-level fields
       • project_custom_attrs (fallback)
       • attachments (PDF/DOCX/plain) 
    2) Try regex on each, in priority order
    3) Deep recursive search as last resort
    """
    logger.info(f"--- Searching Comments for Issue '{issue.get('attributes',{}).get('title','?')}' ---")
    ia = issue.get("attributes", {})
    candidates = []

    # A) Issue custom_attributes
    for c in ia.get("custom_attributes", []):
        term = c.get("term") or c.get("label") or ""
        val  = c.get("value","")
        if isinstance(val, str):
            logger.debug(f"Found custom attribute '{term}'")
            candidates.append(clean_html(val))

    # B) Top-level issue fields
    for fld, val in ia.items():
        if isinstance(val, str) and any(w in fld.lower() for w in ("comment","custom","field")):
            logger.debug(f"Found top-level field '{fld}'")
            candidates.append(clean_html(val))

    # C) Project-level fallback (attributes from parent project)
    for term, val in project_custom_attrs:
        if isinstance(val, str):
            logger.debug(f"Project-level field '{term}' as fallback")
            candidates.append(clean_html(val))

    # D) Attachments
    for name, data in fetch_issue_attachments(issue).items():
        logger.debug(f"Parsing attachment '{name}' for text")
        text = None
        if name.lower().endswith(".pdf"):
            text = convert_pdf_to_text(data)
        else:
            try:
                text = data.decode("utf-8", errors="ignore")
            except:
                text = ""
        if text:
            candidates.append(text)

    # E) Run regex extraction in order
    cm1 = cm2 = ""
    for txt in candidates:
        if not cm1:
            c1, _ = extract_comments_from_text(txt)
            if c1:
                cm1 = c1
                logger.info(f"→ Matched CM1 in candidate: {txt[:80]!r}")
        if not cm2:
            _, c2 = extract_comments_from_text(txt)
            if c2:
                cm2 = c2
                logger.info(f"→ Matched CM2 in candidate: {txt[:80]!r}")
        if cm1 and cm2:
            break

    # F) Deep nested search if still missing
    if not (cm1 and cm2):
        d1, d2 = deep_find_comments(ia)
        if d1 and not cm1:
            cm1 = d1; logger.info("→ Found CM1 via deep recursive search")
        if d2 and not cm2:
            cm2 = d2; logger.info("→ Found CM2 via deep recursive search")

    # G) Final raw-JSON regex fallback
    if not (cm1 and cm2):
        raw = json.dumps(ia)
        if not cm1:
            cm1 = extract_from_text(raw, 1)
            if cm1: logger.info("→ Found CM1 in raw JSON")
        if not cm2:
            cm2 = extract_from_text(raw, 2)
            if cm2: logger.info("→ Found CM2 in raw JSON")

    logger.info(f"Result → CM1={cm1 or '❌'} | CM2={cm2 or '❌'}")
    return cm1, cm2



logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)


def ensure_str(x):
    return str(x) if x is not None else ""


def extract_management_comments(text):
    """
    Try to pull MC1/MC2 from plain text.
    """
    pat1 = re.compile(r"(?:Management|Manager)?\s*Comment\s*1[:\-–]?\s*(.+)", re.IGNORECASE)
    pat2 = re.compile(r"(?:Management|Manager)?\s*Comment\s*2[:\-–]?\s*(.+)", re.IGNORECASE)
    mc1, mc2 = "", ""

    text = text.strip()

    m1 = pat1.search(text)
    if m1:
        mc1 = m1.group(1).strip()

    m2 = pat2.search(text)
    if m2:
        mc2 = m2.group(1).strip()

    return mc1, mc2


def deep_custom_field_search(data):
    """
    Go through ALL nested fields.
    If any field name looks like custom field, grab its value.
    Then also see if its value contains MC1/MC2.
    """
    found_custom = []
    found_mc1 = ""
    found_mc2 = ""

    def walk(x):
        nonlocal found_custom, found_mc1, found_mc2
        if isinstance(x, dict):
            for k, v in x.items():
                k_l = str(k).lower()
                if "custom" in k_l or "field" in k_l:
                    val = BeautifulSoup(str(v), "html.parser").get_text(separator="\n").strip()
                    if val:
                        found_custom.append(val)
                        mc1, mc2 = extract_management_comments(val)
                        if mc1 and not found_mc1:
                            found_mc1 = mc1
                        if mc2 and not found_mc2:
                            found_mc2 = mc2
                walk(v)
        elif isinstance(x, list):
            for i in x:
                walk(i)
        elif isinstance(x, str):
            mc1, mc2 = extract_management_comments(x)
            if mc1 and not found_mc1:
                found_mc1 = mc1
            if mc2 and not found_mc2:
                found_mc2 = mc2

    walk(data)

    return found_custom, found_mc1, found_mc2
# ─── Helpers ────────────────────────────────────────────────
def ensure_str(val):
    """Flatten lists/tuples to comma string, else str()."""
    if isinstance(val, (list, tuple)):
        return ", ".join(map(str, val))
    return str(val or "")

def fetch_attachment_text(url):
    """
    Download an attachment, attempt to extract text:
    - If .pdf → use pdf2docx
    - If .docx → use python-docx
    - Else → raw text.
    """
    logger.debug(f"Fetching attachment: {url}")
    resp = session.get(url, headers=HEADERS)
    resp.raise_for_status()
    content = resp.content
    # Save to temp file by extension
    ext = url.split('.')[-1].lower()
    tmp = f"/tmp/temp_attachment.{ext}"
    with open(tmp, "wb") as f:
        f.write(content)
    if ext == "pdf":
        out = tmp + ".docx"
        Converter(tmp).convert(out, start=0, end=None)
        doc = Document(out)
        text = "\n".join(p.text for p in doc.paragraphs)
    elif ext in ("docx","doc"):
        doc = Document(tmp)
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        text = content.decode(errors="ignore")
    return text

def extract_from_text(raw: str, idx: int) -> str:
    """Final regex fallback over raw JSON or blob."""
    pat = re.compile(
        rf"(?:Management|Manager)?\s*Comment\s*{idx}\s*[:\-–]\s*(.+?)(?=[\"'\n\r]|\}})",
        re.IGNORECASE
    )
    m = pat.search(raw)
    return m.group(1).strip() if m else ""

def find_custom_fields(data: dict):
    """
    Recursively traverse data to collect ALL values of:
    - Custom field(s) 1
    - Custom field(s) 2
    returns (list_of_cf1_texts, list_of_cf2_texts)
    """
    cf1, cf2 = [], []
    key1 = re.compile(r"custom\s*fields?\s*1", re.IGNORECASE)
    key2 = re.compile(r"custom\s*fields?\s*2", re.IGNORECASE)

    def walk(obj):
        if isinstance(obj, dict):
            for k,v in obj.items():
                kl = str(k)
                if key1.search(kl) and isinstance(v, str):
                    cf1.append(BeautifulSoup(v,"html.parser").get_text())
                if key2.search(kl) and isinstance(v, str):
                    cf2.append(BeautifulSoup(v,"html.parser").get_text())
                walk(v)
        elif isinstance(obj, list):
            for item in obj:
                walk(item)
    walk(data)
    return cf1, cf2

def find_management_comments(cf_texts, idx):
    """
    From each custom-field text, pull out embedded
    Management Comment idx or default to whole text.
    """
    pat = re.compile(rf"(?:Management|Manager)?\s*Comment\s*{idx}\s*[:\-–]?\s*(.+)", re.IGNORECASE)
    for txt in cf_texts:
        m = pat.search(txt)
        if m:
            return m.group(1).strip()
    # fallback: return first text block if exists
    return cf_texts[0].strip() if cf_texts else ""

def scan_attachments_for_comments(issue_json, idx):
    """
    Look under relationships.attachments.data
    and fetch each, scanning its text blob.
    """
    comments = []
    rels = issue_json.get("relationships", {}).get("attachments", {}).get("data", [])
    for att in rels:
        url = att.get("links",{}).get("related")
        if not url: 
            continue
        try:
            text = fetch_attachment_text(url)
            c = extract_from_text(text, idx)
            if c:
                comments.append(c)
        except Exception as e:
            logger.warning(f"Could not fetch/parse attachment {url}: {e}")
    return comments


def extract_management_comments(custom_attributes):
    if not custom_attributes:
        return "", ""
    cm = {
        (c.get("term") or "").strip().lower(): c.get("value", "")
        for c in custom_attributes
    }

    mgmt_comment_1 = ""
    mgmt_comment_2 = ""

    # 🔄 Updated patterns allow multiple spaces/underscores/dashes, plural Comments/Fields
    patterns_1 = [
        re.compile(r"management[\s_\-]*comments?[\s_\-]*1"),
        re.compile(r"custom[\s_\-]*fields?[\s_\-]*1"),
        re.compile(r"mgmt[\s_\-]*comments?[\s_\-]*1"),
    ]
    patterns_2 = [
        re.compile(r"management[\s_\-]*comments?[\s_\-]*2"),
        re.compile(r"custom[\s_\-]*fields?[\s_\-]*2"),
        re.compile(r"mgmt[\s_\-]*comments?[\s_\-]*2"),
    ]

    for key, value in cm.items():
        if any(pat.search(key) for pat in patterns_1):
            mgmt_comment_1 = value
            break

    for key, value in cm.items():
        if any(pat.search(key) for pat in patterns_2):
            mgmt_comment_2 = value
            break

    return mgmt_comment_1, mgmt_comment_2


def clean_html(raw):
    """Clean HTML, strip <p>, <br>, other tags, normalize spaces & remove leading/trailing junk."""
    if not raw:
        return ''
    if isinstance(raw, list):
        raw = "\n".join(raw)
    # Replace common tags with newlines
    raw = re.sub(r'</p>', '\n', raw, flags=re.IGNORECASE)
    raw = re.sub(r'<br\s*/?>', '\n', raw, flags=re.IGNORECASE)
    # Remove all other HTML tags
    raw = re.sub(r'<[^>]+>', '', raw)
    # Replace &nbsp; with regular space
    raw = raw.replace('&nbsp;', ' ')
    # For each line, strip leading/trailing spaces and collapse internal multiple spaces
    lines = []
    for line in raw.splitlines():
        line = line.strip()
        # Also normalize multiple spaces inside line to single spaces
        line = re.sub(r'[ \t\u00A0]+', ' ', line)
        lines.append(line)
    # Join cleaned lines, remove any empty lines caused by cleaning
    cleaned = '\n'.join(filter(None, lines))
    return cleaned.strip()


import argparse
import sys
import re
from datetime import datetime

def main():
    parser = argparse.ArgumentParser(
        description="Generate Regional Issues Report from HighBond projects"
    )
    parser.add_argument("--region")
    parser.add_argument("--month")
    parser.add_argument("--severity")
    args, _ = parser.parse_known_args()

    # Prompt and normalize
    rf = (args.region or input("► Region filter (Enter for ALL regions): ")).strip().lower()
    mf = args.month or input("► Month filter (YYYY-MM, Enter to skip): ").strip()
    sf = args.severity or input("► Severity filter (comma-sep, Enter for ALL severities): ").strip()
    sev_set = {s.strip().lower() for s in sf.split(",") if s.strip()} if sf else set()

    # Validate month format if given
    if mf:
        try:
            datetime.strptime(mf, "%Y-%m")
        except ValueError:
            logger.error("Invalid month format. Use YYYY-MM.")
            sys.exit(1)

    projects = get_all_projects()
    projects.sort(key=lambda p: p["attributes"].get("start_date", "" \
    ""), reverse=True)
    logger.info(f"📦 Retrieved {len(projects)} valid projects")

    data = []

    for p in projects:
        attr = p["attributes"]
        project_name = attr.get("name", "")
        start = attr.get("start_date", "")

        if mf and not start.startswith(mf):
            continue

        pid = p["id"]
        ca = attr.get("custom_attributes", [])

        region = ensure_str(next((c["value"] for c in ca if c.get("term") == "Region"), ""))
        if rf and rf != "all" and rf not in region.lower():
            continue

        branch = ensure_str(next((c["value"] for c in ca if c.get("term") == "Branch"), ""))
        bm = ensure_str(next((c["value"] for c in ca if c.get("term") == "Branch Manager"), ""))
        om = ensure_str(next((c["value"] for c in ca if c.get("term") == "Operations Manager"), ""))
        sup = ensure_str(next((c["value"] for c in ca if c.get("term") == "Supervisor"), ""))
        auditors = ensure_str(next((c["value"] for c in ca if c.get("term") == "Auditor(s)"), ""))
        auditor_names = ", ".join(auditors) if auditors else "N/A"

        for isd in get_project_issues(pid):
            ia = isd.get("attributes", {})
            issue_title = ia.get("title", "") or "Untitled Issue"
            sev = ia.get("severity", "").strip().lower()

            # Only filter by severity if user gave any
            if sev_set and sev not in sev_set:
                continue

            cm1, cm2 = "", ""
            cm1_raw, cm2_raw = extract_management_comments(ia.get("custom_attributes", []))
            cm1 = clean_html(cm1_raw)
            cm2 = clean_html(cm2_raw)

            if not cm1 or not cm2:
                proj_cm1_raw, proj_cm2_raw = extract_management_comments(ca)
                if not cm1:
                    cm1 = clean_html(proj_cm1_raw)
                if not cm2:
                    cm2 = clean_html(proj_cm2_raw)

            if not cm1 or not cm2:
                for field in ["description", "effect", "recommendation"]:
                    val = ia.get(field, "")
                    text = BeautifulSoup(str(val), "html.parser").get_text(separator="\n").strip()
                    if text:
                        if not cm1:
                            cm1 = text
                        elif not cm2 and text != cm1:
                            cm2 = text
                    if cm1 and cm2:
                        break

            if not cm1 or not cm2:
                _, deep_cm1, deep_cm2 = deep_custom_field_search(ia)
                if deep_cm1 and not cm1:
                    cm1 = clean_html(deep_cm1)
                if deep_cm2 and not cm2:
                    cm2 = clean_html(deep_cm2)

            cost = ia.get("cost_impact")
            cost = cost if isinstance(cost, (int, float)) else 0.0

            data.append([
                pid,
                project_name,
                branch,
                region,
                start,
                attr.get("status", ""),
                issue_title,
                sev.capitalize(),
                ia.get("description", ""),
                ia.get("effect", ""),
                cost,
                cm1,
                cm2,
                ia.get("recommendation", ""),
                bm,
                om,
                sup,
                auditors,
            ])

    if not data:
        logger.warning("⚠️ No data matched filters.")
        return

    safe_rf = re.sub(r"\W+", "_", rf or "ALL")
    safe_mf = re.sub(r"\W+", "_", mf or "ALL")
    out_fn = f"project_report_{safe_rf}_{safe_mf}.docx"

    doc = create_word_report(data, rf or "ALL", sorted(sev_set) or ["High", "Medium", "Low"])
    doc.save(out_fn)
    logger.info(f"✅ Report saved: {out_fn}")



if __name__ == "_main_":
    try:
        main()
    except ImportError:
        print("Required libraries not found. Please install them using:")
        print("pip install requests pandas beautifulsoup4 python-docx pdf2docx")
        sys.exit(1)
    except Exception as e:
        logger.error(f"An error occurred: {str(e)}")
        logger.exception("Exception details:")
        sys.exit(1)
