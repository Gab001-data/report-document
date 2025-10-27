from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docxtpl import DocxTemplate, InlineImage, Subdoc
import os
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# document = Document()
# section = document.sections[0]
# header = section.header
#print(header.is_linked_to_previous)
# paragraph = header.paragraphs[0]
# paragraph.text = '\t\tHR Audit Report Q1 2025'
# paragraph.style = (document.styles['Header'])

#add new section
# new_section = document.add_section(WD_SECTION.NEW_PAGE)
# new_section.header.paragraphs[0].text = 'Section 2 header'
# new_width, new_height = new_section.page_height, new_section.page_width
# new_section.orientation = WD_ORIENT.LANDSCAPE
# new_section.page_height = new_height
# new_section.page_width = new_width
# document.save('header.docx')

data = {"title": "Monthly KPI", 
           "kpis": [{"name":"Green tea","qty":5, 'revenue':10000},
                    {"name":"Bread",'qty':250, 'revenue':5000},
                    {"name":"Rice",'qty':50, 'revenue':1000}]
          }


print(os.getcwd())          # shows current directory
print(os.path.exists(r"C:\Users\Gabriel O\Desktop\ACL\ACL\Knowledge Base\Python_Reports\py_docx\header_tpl.docx"))

try:
    tpl = DocxTemplate('py_docx/header_tpl.docx')
except Exception as e:
    print(e)
doc = tpl.new_subdoc()
table = doc.add_table(rows=1,cols=3)
#table.style = 'Table Grid'
hdr = table.rows[0].cells # get column title row
hdr[0].text, hdr[1].text, hdr[2].text = tuple(data['kpis'][0].keys())
for item in data["kpis"]:
    data_row = table.add_row()
    cells = data_row.cells
    cells[0].text = item['name']
    cells[1].text = str(item['qty']) if not isinstance(item['qty'],str) else item['qty']
    cells[2].text = str(item['revenue']) if not isinstance(item['revenue'],str) else item['revenue']

    #add cell formatting 
    tcPr = cells[2]._tc.get_or_add_tcPr()
    #shd= tc_pr.add_new_shd()
    if item['revenue'] <= 5000:
        shd_elm = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w'))) # Red shading 
    else:
        shd_elm = parse_xml(r'<w:shd {} w:fill="CCFFCC"/>'.format(nsdecls('w')))
    tcPr.append(shd_elm)

#sd = Subdoc(tpl,doc)

context = {
                    'title': data["title"],
                    'mysubdoc': doc
        }

tpl.render(context)
tpl.save('output/header1.docx')



