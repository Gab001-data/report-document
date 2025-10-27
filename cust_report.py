#import modules 
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#print(help(docx))


# Create a new Document
doc = Document()

#add header of level 0
heading1 = doc.add_heading('Customer Report', level=0)
heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# add sub heading of level 1
subtitle = doc.add_heading('Report prepared by: Gabriel', 2)
subtitle.runs[0].bold = True
subtitle.runs[0].italic = True

#add paragraph
doc_para = doc.add_paragraph('This is a sample customer report generated using Python and the python-docx library.')
doc_para.add_run(' It includes various sections and formatting options.').bold = True
doc_para.add_run(' You can customize it as per your requirements.').italic = True

#add image
doc.add_picture('logo.jpg')

#define a functio to add background color to cells
def set_cell_background_color(cell, color_hex):
    """Set the background color of a cell."""

    # Get the table cell's properties
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    #create a new shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex) #set the fill color
    tcPr.append(shd)

#add table 
# student data 
students = [('Alice', '25', 'New York'), ('Bob', '30', 'San Francisco'), ('Charlie', '22', 'Los Angeles')]
row_count = len(students)
column_count = len(students[0])

table = doc.add_table(rows=1, cols = 3)
header_cells = table.rows[0].cells
header_cells[0].text = 'Name'
header_cells[1].text = 'Age'
header_cells[2].text = 'City'
# for row in range(row_count):
#     for col in range(column_count):
#         cell = table.cell(row+1, col)
#         cell.text = students[row-1][col]
#         if cell.text.isdigit(): 
#             set_cell_background_color(cell, 'FF0000')  # Set background color to yellow for numeric cells
#             cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Set text color to white

for student in students:
    cells = table.add_row().cells
    cells[0].text = student[0]
    cells[1].text = student[1]      
    cells[2].text = student[2]
    # Set background color for numeric cells
    if cells[0].text.isdigit():
        set_cell_background_color(cells[0], 'FF0000')
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    if cells[1].text.isdigit():
        set_cell_background_color(cells[1], 'FF0000')
        cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    if cells[2].text.isdigit():
        set_cell_background_color(cells[2], 'FF0000')
        cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.style = 'Table Grid'
#table.style = 'Grid Table 5 Dark Accent 1'


doc.save('customer_report.docx')
